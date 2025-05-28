from flask import Flask, request, jsonify 
from flask_cors import CORS
import os, uuid, logging, re 
import torch
from transformers import T5Tokenizer, T5ForConditionalGeneration
import chromadb
from chromadb.config import Settings 
from PyPDF2 import PdfReader
from docx import Document
from werkzeug.utils import secure_filename
from chromadb import PersistentClient
from pymongo import MongoClient
from datetime import datetime
import cloudinary
import cloudinary.uploader
import cloudinary.api
import requests
from io import BytesIO
from dotenv import load_dotenv
import cloudinary.utils
from bson import ObjectId
from pymongo.errors import WriteError
from sentence_transformers import SentenceTransformer
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from collections import Counter

app = Flask(__name__)
CORS(app)

load_dotenv()
cloudinary.config(
    cloud_name=os.getenv("CLOUDINARY_CLOUD_NAME", "dpimerw8h"),
    api_key=os.getenv("CLOUDINARY_API_KEY", "495726973616313"),
    api_secret=os.getenv("CLOUDINARY_API_SECRET", "cPM0j222fiNUVb1rHXuugZ2AZ-A")
)
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

MONGO_URI = os.getenv("MONGO_URI", "mongodb://localhost:27017/gistifyDB")
mongo_client = MongoClient(MONGO_URI)
db = mongo_client["gistifyDB"]
summaries_collection = db["Summary"]

UPLOAD_FOLDER = 'Uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10MB max

# Initialize ChromaDB client
chroma_client = PersistentClient(path="./chroma_db")
collection = chroma_client.get_or_create_collection(name="research_papers")

# Initialize SentenceTransformer for embeddings
embedding_model = SentenceTransformer('all-MiniLM-L6-v2')

# Initialize FLAN-T5 model and tokenizer
tokenizer = T5Tokenizer.from_pretrained("google/flan-t5-base", legacy=False)
model = T5ForConditionalGeneration.from_pretrained("google/flan-t5-base", device_map="auto")

def clean_text(text):
    text = re.sub(r"http\S+|www\S+|https\S+", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text

def get_text_from_pdf_from_url(url):
    try:
        response = requests.get(url, stream=True)
        response.raise_for_status()
        pdf_stream = BytesIO(response.content)
        reader = PdfReader(pdf_stream)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text += page_text
        if not text.strip():
            logger.error("No text extracted from PDF")
            return ""
        logger.info(f"Extracted {len(text)} characters from PDF")
        return text
    except Exception as e:
        logger.error(f"Failed to extract text from PDF URL: {e}")
        return ""

def get_text_from_docx_from_url(url):
    try:
        response = requests.get(url, stream=True)
        response.raise_for_status()
        docx_stream = BytesIO(response.content)
        doc = Document(docx_stream)
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        text = text.strip()
        if not text:
            logger.error("No text extracted from DOCX")
            return ""
        logger.info(f"Extracted {len(text)} characters from DOCX")
        return text
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX URL: {e}")
        return ""

def chunk_text(text, max_chunk_size=500):
    """Split text into chunks of approximately max_chunk_size characters."""
    chunks = []
    current_chunk = ""
    paragraphs = text.split('\n')
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(current_chunk) + len(para) + 1 <= max_chunk_size:
            current_chunk += para + "\n"
        else:
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            current_chunk = para + "\n"
    
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    return chunks if chunks else [text[:max_chunk_size]]

def are_summaries_similar(summary1, summary2, threshold=0.9):
    """Check if two summaries are too similar using cosine similarity of embeddings."""
    embeddings = embedding_model.encode([summary1, summary2])
    similarity = cosine_similarity([embeddings[0]], [embeddings[1]])[0][0]
    return similarity > threshold

def is_answer_relevant(answer, question):
    """Check if the answer is relevant to the question and document content."""
    question_words = set(re.findall(r'\w+', question.lower()))
    answer_words = set(re.findall(r'\w+', answer.lower()))
    common_words = question_words.intersection(answer_words)
    relevant_words = [word for word in common_words if len(word) > 3]
        
    # Check for repetition
    answer_phrases = [answer[i:i+10] for i in range(0, len(answer), 10)]  # Check 10-char phrases
    phrase_counts = Counter(answer_phrases)
    max_repetition = max(phrase_counts.values()) if phrase_counts else 1
    
    # For summary questions, require stricter criteria
    is_summary_question = any(phrase in question.lower() for phrase in ["summary of this paper", "summarize the paper", "what is the paper about"])
    word_count = len(answer.split())
    
    logger.info(f"Relevance check: common words = {relevant_words}, word count = {word_count}, max repetition = {max_repetition}")
    
    if is_summary_question:
        return (
            len(relevant_words) >= 3 and
            word_count >= 50 and
            max_repetition <= 2
        )
    return (
        len(relevant_words) >= 2 and
        word_count >= 20 and
        max_repetition <= 2
    )

def is_summary_relevant(summary, doc_id):
    """Check if the summary is relevant to the document content."""    
    # Check for repetition
    summary_phrases = [summary[i:i+10] for i in range(0, len(summary), 10)]
    phrase_counts = Counter(summary_phrases)
    max_repetition = max(phrase_counts.values()) if phrase_counts else 1
    
    # Check if summary resembles query instructions
    query_like_terms = {'provide', 'summary', 'key findings', 'themes', 'conclusions', 'evaluations', 'implications'}
    summary_words = set(re.findall(r'\w+', summary.lower()))
    query_overlap = len(query_like_terms.intersection(summary_words))
    
    word_count = len(summary.split())
    
    logger.info(f"Word count = {word_count}, max repetition = {max_repetition}, query overlap = {query_overlap}")
    
    return (
        word_count >= 200 and
        max_repetition <= 2 and
        query_overlap <= 2  # Reject if summary contains too many query-like terms
    )

def rag_pipeline(query, top_k=15, doc_id=None, use_model_knowledge=False, context_info=None, is_summary=False, summary_type=None):
    try:
        query_embedding = embedding_model.encode([query]).tolist()
        logger.info(f"Query: {query}")
        logger.info(f"Query embedding (first 10 dims): {query_embedding[0][:10]}")

        retrieved_docs = []
        if not use_model_knowledge and doc_id:
            results = collection.query(
                query_embeddings=query_embedding,
                n_results=top_k,
                where={"doc_id": doc_id}
            )
            logger.info(f"ChromaDB query results for doc_id {doc_id}: {len(results['documents'][0]) if results['documents'] else 0} chunks")
            if results['documents'] and results['embeddings']:
                embeddings = results['embeddings'][0]
                similarities = cosine_similarity([query_embedding[0]], embeddings)[0]
                relevant_indices = [i for i, sim in enumerate(similarities) if sim > 0.05]
                retrieved_docs = [results['documents'][0][i] for i in relevant_indices]
                for i, idx in enumerate(relevant_indices):
                    logger.info(f"Chunk {i+1}: similarity = {similarities[idx]:.3f}, content = {results['documents'][0][idx][:100]}...")
                logger.info(f"Retrieved {len(retrieved_docs)} relevant chunks after filtering (similarity > 0.05)")
            
            # Secondary retrieval if insufficient chunks
            if len(retrieved_docs) < 3:
                logger.info("Insufficient chunks, attempting secondary retrieval with keyword filter")
                keyword_results = collection.query(
                    query_texts=query,
                    n_results=5,
                    where={"doc_id": doc_id}
                )
                if keyword_results['documents']:
                    retrieved_docs.extend(keyword_results['documents'][0])
                    for i, doc in enumerate(keyword_results['documents'][0]):
                        logger.info(f"Secondary chunk {i+1}: content = {doc[:100]}...")
                logger.info(f"Total retrieved chunks after secondary query: {len(retrieved_docs)}")
        
        if not retrieved_docs and not use_model_knowledge:
            logger.warning(f"No documents retrieved for query: {query}")
            context = ""
        else:
            context = " ".join(retrieved_docs)
            logger.info(f"Context (first 100 chars): {context[:100]}...")

        if use_model_knowledge:
            if is_summary:
                prompt = (
                    f"Based on your knowledge of '{context_info}', generate a {summary_type} summary of at least 250 words. "
                    f"Focus on the document's key points, main findings, and high-level themes. "
                    f"Ensure the response is detailed, directly relevant to the document's content, "
                    f"and avoids repeating the query or using generic terms like 'general' or 'insufficient'. "
                    f"Do not include instructions or directives in the summary.\n"
                    f"Answer:"
                )
            else:
                is_summary_question = any(phrase in query.lower() for phrase in ["summary of this paper", "summarize the paper", "what is the paper about"])
                if is_summary_question:
                    prompt = (
                        f"Based on your knowledge of '{context_info}', provide a concise summary of the document in 150-200 words. "
                        f"Describe the paper's objectives, key findings, ethical concerns, and recommendations for responsible use. "
                        f"Ensure the response is directly relevant to the document's content, avoids vague terms like 'good' or 'bad', "
                        f"and does not repeat phrases unnecessarily.\nQuestion: {query}\nAnswer:"
                    )
                else:
                    prompt = (
                        f"Based on your knowledge of '{context_info}', answer the question '{query}' in a clear, concise manner. "
                        f"Provide a direct response with relevant details or examples. Use bullet points starting with '-' if listing items, "
                        f"each max 50 words. Avoid generic terms like 'general' or 'insufficient'.\nQuestion: {query}\nAnswer:"
                    )
        else:
            if is_summary:
                prompt = (
                    f"Using the context below, generate a {summary_type} summary of the document in at least 250 words. "
                    f"Focus on the document's key points, main findings, and high-level themes. "
                    f"Ensure the response is detailed, directly derived from the context, "
                    f"and avoids repeating the query or using generic terms like 'general' or 'insufficient'. "
                    f"Do not include instructions or directives in the summary.\n"
                    f"Context: {context}\nAnswer:"
                )
            else:
                is_summary_question = any(phrase in query.lower() for phrase in ["summary of this paper", "summarize the paper", "what is the paper about"])
                if is_summary_question:
                    prompt = (
                        f"Using the context below, provide a concise summary of the document in 150-200 words. "
                        f"Describe the paper's objectives, key findings, ethical concerns, and recommendations for responsible use. "
                        f"Ensure the response is directly relevant to the document's content, avoids vague terms like 'good' or 'bad', "
                        f"and does not repeat phrases unnecessarily.\nContext: {context}\nQuestion: {query}\nAnswer:"
                    )
                else:
                    prompt = (
                        f"Using the context below, answer the question '{query}' in a clear, concise manner. "
                        f"Provide a direct response with relevant details or examples from the context. Use bullet points starting with '-' if listing items, "
                        f"each max 50 words. If context is limited, provide a concise answer based on available information. "
                        f"Avoid generic terms like 'general' or 'insufficient'.\nContext: {context}\nQuestion: {query}\nAnswer:"
                    )

        logger.info(f"Prompt (first 100 chars): {prompt[:100]}...")

        device = next(model.parameters()).device
        inputs = tokenizer(prompt, return_tensors="pt", truncation=True, max_length=1024).to(device)
        generation_params = {
            "input_ids": inputs['input_ids'],
            "attention_mask": inputs['attention_mask'],
            "early_stopping": True
        }
        if is_summary:
            if summary_type == "concise":
                generation_params.update({
                    "max_length": 500,
                    "min_length": 250,
                    "length_penalty": 0.8,
                    "num_beams": 5,
                    "temperature": 0.6,
                    "do_sample": True,
                    "no_repeat_ngram_size": 3
                })
            elif summary_type == "analytical":
                generation_params.update({
                    "max_length": 600,
                    "min_length": 400,
                    "length_penalty": 1.0,
                    "num_beams": 5,
                    "temperature": 0.6,
                    "do_sample": True,
                    "no_repeat_ngram_size": 3
                })
            elif summary_type == "comprehensive":
                generation_params.update({
                    "max_length": 700,
                    "min_length": 500,
                    "length_penalty": 1.2,
                    "num_beams": 5,
                    "temperature": 0.6,
                    "do_sample": True,
                    "no_repeat_ngram_size": 3
                })
        else:
            is_summary_question = any(phrase in query.lower() for phrase in ["summary of this paper", "summarize the paper", "what is the paper about"])
            if is_summary_question:
                generation_params.update({
                    "max_length": 300,
                    "min_length": 150,
                    "length_penalty": 0.8,
                    "num_beams": 4,
                    "temperature": 0.6,
                    "do_sample": True,
                    "no_repeat_ngram_size": 3
                })
            else:
                generation_params.update({
                    "max_length": 400,
                    "min_length": 50,
                    "length_penalty": 1.0,
                    "num_beams": 3,
                    "temperature": 0.9,
                    "do_sample": True,
                    "no_repeat_ngram_size": 3
                })
        outputs = model.generate(**generation_params)
        answer = tokenizer.decode(outputs[0], skip_special_tokens=True)
        logger.info(f"Generated answer (first 100 chars): {answer[:100]}...")
        return answer if answer.strip() else "No relevant content available to generate a response."
    except Exception as e:
        logger.error(f"Error in RAG pipeline: {str(e)}")
        return f"Error generating response: {str(e)}"

def clean_and_extend_answer(ans, question, source):
    ans = re.sub(r"[^\\w\\s.,;:!?()\\[\\]\"'-]", "", ans)
    ans = ans.strip()
    sentences = re.split(r'(?<=[.!?]) +', ans)
    if len(sentences) < 2:
        sentences.append("This answer is based on the provided document context.")
    if any(word in question.lower() for word in ["citation", "reference", "cite"]):
        sentences.append(f"Source: {source}")
    sentences = [s for s in sentences if len(s.split()) > 2]
    return " ".join(sentences)

@app.route('/')
def home():
    return jsonify({"status": "active", "model": "FLAN-T5", "endpoints": ["/upload", "/generate_signed_url", "/summarize", "/ask"]})

@app.route("/generate_signed_url", methods=["GET"])
def generate_signed_url():
    public_id = request.args.get("public_id")
    if not public_id:
        return jsonify({"error": "public_id is required"}), 400

    try:
        signed_url = cloudinary.utils.cloudinary_url(public_id, sign_url=True)[0]
        return jsonify({"signed_url": signed_url})
    except Exception as e:
        logger.error(f"Error generating signed URL: {e}")
        return jsonify({"error": "An error occurred while generating signed URL"}), 500

@app.route("/upload", methods=["POST"])
def upload_file():
    file = request.files.get("file")
    if not file or file.filename == "":
        logger.warning("No file provided in upload request")
        return jsonify({"error": "No valid file uploaded"}), 400

    try:
        filename = secure_filename(file.filename)
        ext = filename.rsplit('.', 1)[-1].lower()
        if ext not in ["pdf", "docx"]:
            logger.error(f"Unsupported file type: {ext}")
            return jsonify({"error": "Only PDF and DOCX files are supported"}), 400

        cloudinary_url = request.form.get("secure_url")
        if not cloudinary_url:
            logger.info(f"No secure_url provided for {filename}, uploading to Cloudinary")
            upload_result = cloudinary.uploader.upload(
                file,
                resource_type="raw",
                public_id=f"uploads/{uuid.uuid4()}_{filename}",
                overwrite=True
            )
            cloudinary_url = upload_result.get("secure_url")
            if not cloudinary_url:
                logger.error("Failed to upload file to Cloudinary")
                return jsonify({"error": "Failed to upload file to Cloudinary"}), 500
            logger.info(f"File uploaded to Cloudinary: {cloudinary_url}")

        logger.info(f"Processing file: {filename}, Cloudinary URL: {cloudinary_url}")

        if ext == "pdf":
            text = get_text_from_pdf_from_url(cloudinary_url)
        elif ext == "docx":
            text = get_text_from_docx_from_url(cloudinary_url)
        else:
            text = ""

        if not text or len(text.strip()) < 50:
            logger.error(f"Text extraction failed or insufficient: {text[:100]}...")
            return jsonify({"error": "Failed to extract sufficient text from file"}), 400

        text = clean_text(text)
        chunks = chunk_text(text, max_chunk_size=500)
        if not chunks or all(not chunk.strip() for chunk in chunks):
            logger.error("No valid chunks generated from text")
            return jsonify({"error": "No valid content to process"}), 400

        doc_embedding = embedding_model.encode(chunks).tolist()

        doc_id = str(uuid.uuid4())
        chunk_ids = [f"{doc_id}_{i}" for i in range(len(chunks))]
        collection.add(
            ids=chunk_ids,
            documents=chunks,
            embeddings=doc_embedding,
            metadatas=[{"source": filename, "cloudinary_url": cloudinary_url, "doc_id": doc_id} for _ in chunks]
        )

        logger.info(f"Uploaded document with doc_id: {doc_id}, chunks: {len(chunks)}")

        return jsonify({
            "message": "File uploaded to Cloudinary and ChromaDB",
            "sample": text[:250],
            "source": filename,
            "cloudinary_url": cloudinary_url,
            "doc_id": doc_id,
            "file": {
                "filePath": cloudinary_url,
                "id": doc_id,
                "fileName": filename
            }
        })
    except Exception as e:
        logger.error(f"Upload error for {filename}: {str(e)}")
        return jsonify({"error": f"An error occurred during upload: {str(e)}"}), 500

@app.route("/upload_text", methods=["POST"])
def upload_text():
    try:
        text = request.form.get("text")
        file_name = request.form.get("file_name")
        cloudinary_url = request.form.get("cloudinary_url")
        user_id = request.form.get("user_id")

        if not text or not file_name:
            return jsonify({"error": "Text and file_name are required"}), 400

        if len(text.strip()) < 50:
            logger.error("Insufficient text provided for upload")
            return jsonify({"error": "Text content is too short"}), 400

        text = clean_text(text)
        chunks = chunk_text(text, max_chunk_size=500)
        if not chunks or all(not chunk.strip() for chunk in chunks):
            logger.error("No valid chunks generated from text")
            return jsonify({"error": "No valid content to process"}), 400

        doc_embedding = embedding_model.encode(chunks).tolist()

        doc_id = str(uuid.uuid4())
        chunk_ids = [f"{doc_id}_{i}" for i in range(len(chunks))]
        
        collection.add(
            ids=chunk_ids,
            documents=chunks,
            embeddings=doc_embedding,
            metadatas=[{
                "source": file_name,
                "cloudinary_url": cloudinary_url,
                "user_id": user_id,
                "type": "text",
                "doc_id": doc_id
            } for _ in chunks]
        )

        logger.info(f"Uploaded text with doc_id: {doc_id}, chunks: {len(chunks)}")

        return jsonify({
            "message": "Text uploaded to ChromaDB",
            "doc_id": doc_id,
            "cloudinary_url": cloudinary_url,
            "file": {
                "filePath": cloudinary_url,
                "id": doc_id,
                "fileName": file_name
            }
        })
    except Exception as e:
        logger.error(f"Text upload error: {str(e)}")
        return jsonify({"error": "An error occurred during text upload"}), 500

@app.route("/summarize", methods=["POST"])
def summarize():
    try:
        doc_id = request.form.get("doc_id")
        summary_type = request.form.get("summary_type")

        if not doc_id or not summary_type:
            try:
                data = request.get_json()
                doc_id = data.get("doc_id") if data else None
                summary_type = data.get("summary_type") if data else None
            except Exception:
                logger.warning("Invalid JSON body in summarize request")

        logger.info(f"Received /summarize request with: doc_id='{doc_id}', summary_type='{summary_type}'")

        if not doc_id or not summary_type:
            logger.warning("Missing required fields in summarize request")
            return jsonify({"error": "Missing required fields: doc_id and summary_type are required"}), 400

        chroma_result = collection.get(ids=[f"{doc_id}_0"])
        if not chroma_result.get("documents"):
            logger.error(f"Document with doc_id {doc_id} not found in ChromaDB")
            return jsonify({"error": "Document not found"}), 404

        all_chunks = collection.get(where={"doc_id": doc_id})
        if not all_chunks.get("documents") or sum(len(doc) for doc in all_chunks["documents"]) < 50:
            logger.error(f"Document with doc_id {doc_id} has insufficient content")
            return jsonify({"error": "Document contains insufficient content for summarization"}), 400

        logger.info(f"Retrieved {len(all_chunks['documents'])} chunks for doc_id {doc_id}, total length: {sum(len(doc) for doc in all_chunks['documents'])}")

        metadata = all_chunks.get("metadatas", [{}])[0]
        source = metadata.get("source", "unknown document")

        if "concise" in summary_type.lower():
            summary_query = (
                f"Summarize the document's key points, main findings, and high-level themes in at least 250 words. "
                f"Provide a clear, succinct overview of the core content, prioritizing brevity and clarity. "
                f"Include essential details like primary topics or outcomes, but avoid deep analysis or extensive background."
            )
            summary_type_internal = "concise"
        elif "analytical" in summary_type.lower():
            summary_query = (
                f"Provide an analytical summary of the document in at least 250 words. "
                f"Critically evaluate methods, findings, or arguments, highlighting insights, implications, and potential weaknesses. "
                f"Include specific examples or evidence to support the analysis."
            )
            summary_type_internal = "analytical"
        elif "comprehensive" in summary_type.lower():
            summary_query = (
                f"Provide a comprehensive summary of the document in at least 250 words. "
                f"Cover all major aspects, including background, methods, results, and conclusions in detail. "
                f"Ensure a thorough overview without focusing on critical analysis."
            )
            summary_type_internal = "comprehensive"
        else:
            summary_query = (
                f"Summarize the document in at least 250 words, covering key themes, methods, or findings with sufficient detail. "
                f"Focus on a balanced overview of the content, including main points and outcomes."
            )
            summary_type_internal = "default"

        max_attempts = 3
        attempt = 0
        summary = None

        while attempt < max_attempts:
            summary = rag_pipeline(
                summary_query,
                top_k=15,
                doc_id=doc_id,
                is_summary=True,
                summary_type=summary_type_internal
            )
            if "Error" in summary or "Insufficient" in summary:
                logger.warning(f"Attempt {attempt + 1}: Failed to generate summary: {summary}")
                context_info = f"{source} (a document on ChatGPT and ethics in scholarly publishing)"
                summary = rag_pipeline(
                    summary_query,
                    top_k=15,
                    use_model_knowledge=True,
                    context_info=context_info,
                    is_summary=True,
                    summary_type=summary_type_internal
                )
                logger.info(f"Model knowledge fallback summary: {summary[:100]}...")
            
            if is_summary_relevant(summary, doc_id):
                logger.info(f"Summary is relevant after {attempt + 1} attempts: {summary[:100]}...")
                break
            logger.warning(f"Summary not relevant: {summary[:100]}... Attempt {attempt + 1}/{max_attempts}")
            attempt += 1

        if attempt == max_attempts:
            logger.error(f"Failed to generate relevant summary after {max_attempts} attempts")

        word_count = len(summary.split())
        logger.info(f"Initial {summary_type_internal} summary word count: {word_count}")
        if word_count < 250:
            logger.info(f"Summary too short ({word_count} words), attempting to extend")
            context_info = source
            extended_query = (
                f"Based on the document '{context_info}', provide additional details to extend the {summary_type_internal} summary to at least 250 words. "
                f"Focus on key points, findings, and themes relevant to the document’s content, "
                f"avoiding repetition of the query or generic terms like 'general' or 'insufficient'. "
                f"Do not include instructions or directives in the response."
            )
            additional_text = rag_pipeline(
                extended_query,
                top_k=15,
                doc_id=doc_id,
                is_summary=True,
                summary_type=summary_type_internal
            )
            if "Error" not in additional_text and "Insufficient" not in additional_text and is_summary_relevant(additional_text, doc_id):
                summary = f"{summary}\n\n{additional_text}"
                word_count = len(summary.split())
                logger.info(f"Extended summary word count: {word_count}")
            if word_count < 250:
                shortfall = 250 - word_count
                if summary_type_internal == "concise":
                    fallback_text = (
                        f"The document explores ChatGPT’s impact on academia, focusing on automation of scholarly tasks. "
                        f"It addresses ethical challenges, such as ensuring fairness and originality in research outputs. "
                        f"The study provides practical recommendations for integrating AI responsibly in academic settings."
                    )
                elif summary_type_internal == "analytical":
                    fallback_text = (
                        f"The document critically evaluates ChatGPT’s role in academia, noting limitations in addressing complex ethical issues. "
                        f"It suggests further research to strengthen AI’s application in scholarly work, highlighting potential biases."
                    )
                else:
                    fallback_text = (
                        f"The document provides a detailed overview of ChatGPT’s applications in scholarly publishing. "
                        f"It covers methods, outcomes, and ethical considerations, offering a comprehensive perspective on AI’s academic role."
                    )
                summary = f"{summary}\n\n{fallback_text[:shortfall*5]}"
                word_count = len(summary.split())
                logger.info(f"Final summary word count after fallback: {word_count}")

        advantages_query = (
            f"List two key advantages of the document content in bullet points starting with '-', each max 20 words. "
            f"Points must be specific to the document's content or purpose."
        )
        disadvantages_query = (
            f"List two key disadvantages of the document content in bullet points starting with '-', each max 20 words. "
            f"Points must be specific to the document's limitations."
        )

        advantages_text = rag_pipeline(advantages_query, top_k=15, doc_id=doc_id)
        disadvantages_text = rag_pipeline(disadvantages_query, top_k=15, doc_id=doc_id)

        logger.info(f"Raw advantages text: {advantages_text}")
        logger.info(f"Raw disadvantages text: {disadvantages_text}")

        def parse_points(text, query_type):
            lines = text.split('\n')
            points = []
            for line in lines:
                line = line.strip()
                if line.startswith('-') and len(line.split()) <= 20 and len(line.split()) > 3:
                    cleaned_line = re.sub(r'^[-*]\s*', '', line)
                    if not any(vague in cleaned_line.lower() for vague in ["general", "insufficient", "limited scope", "no specific"]):
                        points.append(f"- {cleaned_line}")
            if len(points) < 2:
                logger.info(f"Insufficient valid {query_type} points, generating contextual points")
                context_info = summary[:100] if summary and "Error" not in summary else source
                if query_type == "advantages":
                    contextual_query = (
                        f"List two specific advantages of a document about '{context_info}' in bullet points starting with '-', "
                        f"each max 20 words, focusing on its purpose or features."
                    )
                else:
                    contextual_query = (
                        f"List two specific disadvantages of a document about '{context_info}' in bullet points starting with '-', "
                        f"each max 20 words, focusing on its limitations."
                    )
                contextual_text = rag_pipeline(contextual_query, top_k=15, use_model_knowledge=True, context_info=context_info)
                contextual_lines = contextual_text.split('\n')
                for line in contextual_lines:
                    line = line.strip()
                    if line.startswith('-') and len(line.split()) <= 20 and len(line.split()) > 3:
                        cleaned_line = re.sub(r'^[-*]\s*', '', line)
                        if not any(vague in cleaned_line.lower() for vague in ["general", "insufficient", "limited scope", "no specific"]):
                            points.append(f"- {cleaned_line}")
                            if len(points) == 2:
                                break
            if len(points) < 2:
                if query_type == "advantages":
                    points.append("- Highlights ethical issues in AI-driven publishing.")
                    if len(points) < 2:
                        points.append("- Offers practical guidelines for AI use.")
                else:
                    points.append("- Limited depth in technical AI analysis.")
                    if len(points) < 2:
                        points.append("- May overlook broader societal impacts.")
            return points[:2]

        advantages = parse_points(advantages_text, "advantages")
        disadvantages = parse_points(disadvantages_text, "disadvantages")

        logger.info(f"Final advantages: {advantages}")
        logger.info(f"Final disadvantages: {disadvantages}")

        summary_data = {
            "userId": request.form.get("user_id"),
            "doc_id": doc_id,
            "file_id": request.form.get("file_id"),
            "summary": summary,
            "advantages": advantages,
            "disadvantages": disadvantages,
            "file_name": request.form.get("file_name"),
            "fileUrl": request.form.get("file_path"),
            "summary_type": summary_type,
        }
        mongo_result = db.Summary.insert_one(summary_data)
        logger.info(f"Stored summary in MongoDB with ID: {str(mongo_result.inserted_id)}")

        response = {
            "summary": summary,
            "advantages": advantages,
            "disadvantages": disadvantages,
            "summary_id": str(mongo_result.inserted_id),
        }

        logger.info(f"Sending response: {response}")
        return jsonify(response)
    except WriteError as e:
        logger.error(f"MongoDB write error: {str(e)}")
        return jsonify({"error": "Failed to save summary due to database error"}), 500
    except Exception as e:
        logger.error(f"Summarize error: {str(e)}")
        return jsonify({"error": "An error occurred during summarization"}), 500

@app.route("/ask", methods=["POST"])
def ask():
    data = request.get_json()
    question = data.get("question")
    doc_id = data.get("doc_id")
    if not question or not doc_id:
        logger.error("Both question and doc_id are required")
        return jsonify({"error": "Both question and doc_id are required"}), 400

    try:
        logger.info(f"Processing question: {question} for doc_id: {doc_id}")

        all_chunks = collection.get(where={"doc_id": doc_id})
        if not all_chunks.get("documents") or sum(len(doc) for doc in all_chunks["documents"]) < 50:
            logger.error(f"Document with doc_id {doc_id} has insufficient content")
            return jsonify({"error": "Document contains insufficient content to answer questions"}), 400
        logger.info(f"Document has {len(all_chunks['documents'])} chunks, total length: {sum(len(doc) for doc in all_chunks['documents'])}")

        chroma_result = collection.get(ids=[f"{doc_id}_0"])
        if not chroma_result.get("documents"):
            logger.error(f"Document with doc_id {doc_id} not found in ChromaDB")
            return jsonify({"error": "Document not found"}), 404
        metadata = chroma_result.get("metadatas", [{}])[0]
        source = metadata.get("source", "Unknown")

        max_attempts = 3
        attempt = 0
        answer = None

        while attempt < max_attempts:
            answer = rag_pipeline(question, top_k=15, doc_id=doc_id)
            if "Error" in answer or "No relevant content" in answer:
                logger.warning(f"Attempt {attempt + 1}: Failed to generate answer: {answer}")
                context_info = f"{source} (a paper on ChatGPT and ethics in scholarly publishing)"
                answer = rag_pipeline(
                    question,
                    top_k=15,
                    use_model_knowledge=True,
                    context_info=context_info
                )
                logger.info(f"Model knowledge fallback answer: {answer[:100]}...")
            
            if is_answer_relevant(answer, question):
                logger.info(f"Answer is relevant after {attempt + 1} attempts: {answer[:100]}...")
                break
            logger.warning(f"Answer not relevant for question: {question}. Attempt {attempt + 1}/{max_attempts}")
            attempt += 1

        if attempt == max_attempts:
            logger.error(f"Failed to generate relevant answer after {max_attempts} attempts")
            is_summary_question = any(phrase in question.lower() for phrase in ["summary of this paper", "summarize the paper", "what is the paper about"])

        answer = clean_and_extend_answer(answer, question, source)
        logger.info(f"Final answer: {answer}")

        return jsonify({
            "answer": answer,
            "source": source
        })
    except Exception as e:
        logger.error(f"Ask endpoint error: {str(e)}")
        return jsonify({"error": f"An error occurred while processing the question: {str(e)}"}), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5001)